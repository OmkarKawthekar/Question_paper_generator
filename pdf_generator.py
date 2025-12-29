"""
Document generation helpers:
- build_question_paper_pdf(sections) -> bytes
- build_question_paper_docx(sections) -> bytes
- generate_sample_syllabus_pdf_bytes() -> bytes (for convenience)

Sections format expected by the app:
  [
    ("Section: 4 Mark Questions", [ {"unit", "question", "marks", "difficulty"}, ... ]),
    ("Section: 6 Mark Questions", [ ... ])
  ]
"""

from __future__ import annotations

import io
from typing import List, Tuple, Dict

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.styles import ParagraphStyle

from docx import Document
from docx.shared import Pt


def build_question_paper_pdf(sections: List[Tuple[str, List[Dict]]]) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='TitleCenter', parent=styles['Title'], alignment=TA_CENTER)
    section_style = styles['Heading2']
    question_style = styles['BodyText']

    elements = []
    elements.append(Paragraph("Question Paper", title_style))
    elements.append(Spacer(1, 12))

    for section_title, questions in sections:
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(section_title, section_style))
        elements.append(Spacer(1, 6))
        for idx, q in enumerate(questions, start=1):
            q_text = f"{idx}. [{q['marks']}M] ({q['difficulty']}) {q['question']} — <i>{q['unit']}</i>"
            elements.append(Paragraph(q_text, question_style))
            elements.append(Spacer(1, 4))

    doc.build(elements)
    return buffer.getvalue()


def build_question_paper_docx(sections: List[Tuple[str, List[Dict]]]) -> bytes:
    doc = Document()
    title = doc.add_heading("Question Paper", level=1)
    title.alignment = 1

    for section_title, questions in sections:
        doc.add_heading(section_title, level=2)
        for idx, q in enumerate(questions, start=1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. [{q['marks']}M] ({q['difficulty']}) {q['question']} — {q['unit']}")
            run.font.size = Pt(11)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def generate_sample_syllabus_pdf_bytes() -> bytes:
    """Create a small syllabus PDF with simple 'Unit' headings for demo."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    normal = styles['BodyText']
    h = styles['Heading2']

    elements = []
    elements.append(Paragraph("Sample Syllabus", styles['Title']))

    elements.append(Paragraph("Unit 1", h))
    elements.append(Paragraph("Introduction to Data Structures: arrays, linked lists, stacks, queues; basic operations and applications.", normal))

    elements.append(Paragraph("Unit 2", h))
    elements.append(Paragraph("Trees and Graphs: binary trees, BSTs, traversals, graph representations, BFS/DFS, shortest paths.", normal))

    elements.append(Paragraph("Unit 3", h))
    elements.append(Paragraph("Algorithms: sorting, searching, time and space complexity, greedy and dynamic programming basics.", normal))

    doc.build(elements)
    return buffer.getvalue()



